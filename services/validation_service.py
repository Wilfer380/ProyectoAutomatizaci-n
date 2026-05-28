from __future__ import annotations

from pathlib import Path
import tempfile

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

from utils.constants import (
    EXCEL_SHEET_LABEL,
    EXCEL_SHEET_SOURCE,
    SOURCE_HEADERS,
    TARGET_PRINTER_NAME,
)
from utils.normalization import normalize_excel_scalar

try:
    import win32print
except ImportError:  # pragma: no cover
    win32print = None


class ValidationError(Exception):
    pass


class ValidationService:
    def _validate_file_access(self, path: Path, *, readable: bool = True) -> None:
        try:
            if readable:
                with path.open("rb"):
                    pass
            else:
                with path.open("r+b"):
                    pass
        except PermissionError as exc:
            raise ValidationError(f"No hay permisos suficientes para acceder a '{path}'.") from exc
        except OSError as exc:
            raise ValidationError(f"El archivo '{path}' no se puede abrir en este momento. Puede estar bloqueado por otra aplicación.") from exc

    def validate_excel_file(self, excel_path: str) -> Path:
        path = Path(excel_path)
        if not path.exists():
            raise ValidationError("El archivo Excel no existe.")
        if path.suffix.lower() not in {".xlsx", ".xlsm", ".xls"}:
            raise ValidationError("El archivo Excel no tiene una extensión válida.")
        self._validate_file_access(path)
        return path

    def validate_word_file(self, word_path: str) -> Path:
        path = Path(word_path)
        if not path.exists():
            raise ValidationError("La plantilla Word no existe.")
        if path.suffix.lower() not in {".docx", ".docm", ".doc"}:
            raise ValidationError("La plantilla Word no tiene una extensión válida.")
        self._validate_file_access(path)
        return path

    def validate_directory_writable(self, directory: str | Path) -> Path:
        path = Path(directory)
        path.mkdir(parents=True, exist_ok=True)
        try:
            with tempfile.NamedTemporaryFile(prefix="automatizacion_sap_check_", dir=str(path), delete=True):
                pass
        except PermissionError as exc:
            raise ValidationError(f"No se puede escribir en la carpeta '{path}'.") from exc
        except OSError as exc:
            raise ValidationError(f"No se pudo validar la carpeta de trabajo '{path}'.") from exc
        return path

    def validate_selected_filter(self, selected_filter: str) -> str:
        value = selected_filter.strip()
        if not value:
            raise ValidationError("Debe seleccionar un filtro antes de comenzar.")
        return value

    def validate_required_sheets(self, sheet_names: list[str]) -> None:
        missing = [name for name in (EXCEL_SHEET_SOURCE, EXCEL_SHEET_LABEL) if name not in sheet_names]
        if missing:
            raise ValidationError(f"No se encontraron las hojas requeridas: {', '.join(missing)}.")

    def validate_required_headers(self, headers: list[str]) -> None:
        missing = [header for header in SOURCE_HEADERS.values() if header not in headers]
        if missing:
            raise ValidationError(f"No se encontraron las columnas requeridas: {', '.join(missing)}.")

    def validate_word_template_placeholders(self, word_path: str | Path, expected_count: int = 27) -> None:
        path = Path(word_path)
        if Document is None:
            raise ValidationError("No se pudo inspeccionar la plantilla Word porque python-docx no está instalado.")

        try:
            doc = Document(str(path))
        except Exception as exc:
            raise ValidationError(f"No se pudo abrir la plantilla Word '{path}' para validar sus placeholders.") from exc

        xml_text = self._collect_docx_text(doc)

        missing = [f"<img{index}>" for index in range(1, expected_count + 1) if f"<img{index}>" not in xml_text]
        if missing:
            raise ValidationError(
                "La plantilla Word no contiene los placeholders requeridos: " + ", ".join(missing)
            )

    @staticmethod
    def _collect_docx_text(doc) -> str:
        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(paragraph.text for paragraph in cell.paragraphs)
        for section in doc.sections:
            parts.extend(paragraph.text for paragraph in section.header.paragraphs)
            parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
        return "\n".join(parts)

    def validate_filter_records(self, records_count: int, selected_filter: str) -> None:
        if records_count <= 0:
            raise ValidationError(f"El filtro seleccionado '{selected_filter}' no tiene registros.")

    def validate_block_size(self, records_count: int, block_size: int) -> None:
        if records_count > block_size:
            raise ValidationError(f"El bloque actual supera el máximo permitido de {block_size} registros.")

    def validate_assets_match(self, source_assets: list[str], generated_assets: list[str]) -> None:
        normalized_source = [normalize_excel_scalar(asset) for asset in source_assets]
        normalized_generated = [normalize_excel_scalar(asset) for asset in generated_assets]

        if normalized_source == normalized_generated:
            return

        mismatches: list[str] = []
        for index, (source, generated, source_norm, generated_norm) in enumerate(
            zip(source_assets, generated_assets, normalized_source, normalized_generated),
            start=1,
        ):
            if source_norm != generated_norm:
                mismatches.append(
                    f"Fila {index}: origen='{source}' normalizado='{source_norm}' salida='{generated}' normalizado='{generated_norm}'"
                )

        if len(normalized_source) != len(normalized_generated):
            mismatches.append(
                f"Cantidad distinta de activos: origen={len(normalized_source)} salida={len(normalized_generated)}"
            )

        raise ValidationError(
            "La validación de activos falló. Diferencias encontradas: " + "; ".join(mismatches)
        )

    def validate_printer_installed(self, printer_name: str = TARGET_PRINTER_NAME) -> None:
        if win32print is None:
            raise ValidationError("No se pudo importar win32print para validar impresoras.")

        installed_printers = [
            info[2]
            for info in win32print.EnumPrinters(
                win32print.PRINTER_ENUM_LOCAL | win32print.PRINTER_ENUM_CONNECTIONS
            )
        ]
        if printer_name not in installed_printers:
            raise ValidationError(f"La impresora '{printer_name}' no está instalada.")
