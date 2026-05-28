from __future__ import annotations

import os
import re
from dataclasses import dataclass
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.text.paragraph import Paragraph
    from docx.shared import Cm
except ImportError:  # pragma: no cover
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Paragraph = None
    Cm = None

from utils.constants import WORD_LABEL_HEIGHT_CM, WORD_LABEL_WIDTH_CM

pythoncom = None
win32 = None


def _load_pywin32():
    global pythoncom, win32
    if pythoncom is not None and win32 is not None:
        return pythoncom, win32

    from utils.win32_bootstrap import load_pywin32

    pythoncom, win32 = load_pywin32()
    return pythoncom, win32


@dataclass(slots=True)
class VisualValidationResult:
    position: int
    slot_name: str
    expected_width_cm: float
    expected_height_cm: float
    applied_width_cm: float
    applied_height_cm: float
    container_width_cm: float | None
    container_height_cm: float | None
    adjusted: bool
    details: str
    shape_name: str
    page_number: int
    created_slot: bool


@dataclass(slots=True)
class BlankPageCleanupResult:
    pages_before: int
    pages_after: int
    removed_pages: int
    removed_page_numbers: list[int]
    cleaned_placeholders: list[str]
    remaining_placeholders: list[str]
    expected_last_image_page: int | None = None
    validation_passes: int = 0


@dataclass(slots=True)
class ImageCountValidationResult:
    expected_count: int
    detected_count: int
    missing_count: int
    extra_count: int
    missing_positions: list[int]
    repaired_positions: list[int]


class WordService:
    MAX_PLACEHOLDER_IMAGES = 27

    def __init__(self) -> None:
        self.document = None
        self.document_path: Path | None = None
        self.last_layout_details = "plantilla Word no inspeccionada"
        self._embedded_image_count = 0
        self._review_app = None
        self._review_document = None
        self._review_com_initialized = False

    def open(self, template_path: str, visible: bool = False) -> None:  # noqa: ARG002
        self._ensure_docx_support()
        path = Path(template_path).resolve()
        self.document = Document(str(path))
        self.document_path = path
        self._embedded_image_count = 0

    def create_from_template(self, template_path: str, visible: bool = False) -> None:
        self.open(template_path, visible=visible)

    def close(self, save_changes: bool = False) -> None:  # noqa: ARG002
        self.document = None
        self.document_path = None
        self._embedded_image_count = 0
        self.close_review_document(save_changes=save_changes)

    def prepare_placeholder_document(self) -> None:
        if self.document is None:
            raise RuntimeError("Word no está abierto.")
        remaining = self.find_remaining_image_placeholders()
        self.last_layout_details = (
            "plantilla Word lista para reemplazar placeholders "
            f"<img1>...<img{self.MAX_PLACEHOLDER_IMAGES}>; pendientes={len(remaining)}"
        )

    def get_label_layout_details(self) -> str:
        return self.last_layout_details

    def replace_image_placeholder(self, placeholder_number: int, image_path: str | Path) -> VisualValidationResult:
        if self.document is None:
            raise RuntimeError("Word no está abierto.")

        placeholder = self._placeholder_text(placeholder_number)
        image_file = Path(image_path)
        image_inserted = False
        placeholder_found = False

        for paragraph in self._iter_docx_paragraphs(self.document):
            if placeholder not in paragraph.text:
                continue
            placeholder_found = True
            image_inserted = self._replace_placeholder_in_paragraph(
                paragraph,
                placeholder=placeholder,
                image_path=image_file if image_file.exists() else None,
            ) or image_inserted

        if not placeholder_found:
            return VisualValidationResult(
                position=placeholder_number,
                slot_name=placeholder,
                expected_width_cm=WORD_LABEL_WIDTH_CM,
                expected_height_cm=WORD_LABEL_HEIGHT_CM,
                applied_width_cm=0.0,
                applied_height_cm=0.0,
                container_width_cm=None,
                container_height_cm=None,
                adjusted=False,
                details=f"No se encontró el placeholder {placeholder}.",
                shape_name="",
                page_number=1,
                created_slot=False,
            )

        if not image_file.exists():
            return VisualValidationResult(
                position=placeholder_number,
                slot_name=placeholder,
                expected_width_cm=WORD_LABEL_WIDTH_CM,
                expected_height_cm=WORD_LABEL_HEIGHT_CM,
                applied_width_cm=0.0,
                applied_height_cm=0.0,
                container_width_cm=None,
                container_height_cm=None,
                adjusted=False,
                details=f"imagen faltante: {image_file}; placeholder limpiado",
                shape_name="",
                page_number=1,
                created_slot=False,
            )

        if image_inserted:
            self._embedded_image_count += 1

        return VisualValidationResult(
            position=placeholder_number,
            slot_name=placeholder,
            expected_width_cm=WORD_LABEL_WIDTH_CM,
            expected_height_cm=WORD_LABEL_HEIGHT_CM,
            applied_width_cm=WORD_LABEL_WIDTH_CM if image_inserted else 0.0,
            applied_height_cm=WORD_LABEL_HEIGHT_CM if image_inserted else 0.0,
            container_width_cm=None,
            container_height_cm=None,
            adjusted=image_inserted,
            details=f"placeholder {placeholder} reemplazado por imagen" if image_inserted else f"placeholder {placeholder} limpiado",
            shape_name="",
            page_number=1,
            created_slot=False,
        )

    def clear_unused_image_placeholders(self, used_count: int, max_count: int | None = None) -> list[str]:
        if self.document is None:
            raise RuntimeError("Word no está abierto.")

        max_count = max_count or self.MAX_PLACEHOLDER_IMAGES
        cleaned: list[str] = []
        for placeholder_number in range(used_count + 1, max_count + 1):
            placeholder = self._placeholder_text(placeholder_number)
            if self._replace_all_exact_text(placeholder, ""):
                cleaned.append(placeholder)
        return cleaned

    def validate_embedded_image_count(self, expected_count: int) -> ImageCountValidationResult:
        detected_count = self._embedded_image_count
        missing_count = max(expected_count - detected_count, 0)
        extra_count = max(detected_count - expected_count, 0)
        missing_positions = list(range(detected_count + 1, expected_count + 1)) if missing_count else []
        return ImageCountValidationResult(expected_count, detected_count, missing_count, extra_count, missing_positions, [])

    def count_visible_images(self) -> int:
        return self._embedded_image_count

    def cleanup_blank_pages(
        self,
        expected_last_image_page: int | None = None,
        validation_passes: int = 3,
        cleaned_placeholders: list[str] | None = None,
    ) -> BlankPageCleanupResult:
        remaining = self.find_remaining_image_placeholders() if self.document is not None else []
        return BlankPageCleanupResult(
            pages_before=1,
            pages_after=1,
            removed_pages=0,
            removed_page_numbers=[],
            cleaned_placeholders=cleaned_placeholders or [],
            remaining_placeholders=remaining,
            expected_last_image_page=expected_last_image_page,
            validation_passes=validation_passes,
        )

    def find_remaining_image_placeholders(self) -> list[str]:
        if self.document is None:
            raise RuntimeError("Word no está abierto.")
        remaining: list[str] = []
        for placeholder_number in range(1, self.MAX_PLACEHOLDER_IMAGES + 1):
            placeholder = self._placeholder_text(placeholder_number)
            if self._contains_exact_text(placeholder):
                remaining.append(placeholder)
        return remaining

    def get_page_count(self) -> int:
        return 1

    def print_document(self, printer_name: str) -> None:  # noqa: ARG002
        raise RuntimeError("La impresión automática fue deshabilitada. Validá e imprimí el documento manualmente.")

    def save_document_copy(self, output_path: str) -> None:
        if self.document is None:
            raise RuntimeError("Word no está abierto.")
        output_file = Path(output_path).resolve()
        output_file.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(output_file))
        self.document_path = output_file

    def show_to_user(self) -> None:
        if self.document_path is None:
            raise RuntimeError("No hay documento Word generado para abrir.")
        os.startfile(str(self.document_path))

    def open_for_review(self, document_path: str) -> None:
        pythoncom_module, win32_module = _load_pywin32()
        self.close_review_document(save_changes=False)

        pythoncom_module.CoInitialize()
        self._review_com_initialized = True
        target = str(Path(document_path).resolve())
        self._review_app = win32_module.dynamic.Dispatch("Word.Application")
        self._review_app.Visible = True
        self._review_app.DisplayAlerts = 0
        self._review_document = self._review_app.Documents.Open(
            target,
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
            Visible=True,
        )
        self._review_document.Activate()
        try:
            self._review_app.Options.PrintDrawingObjects = True
        except Exception:
            pass
        try:
            view = self._review_app.ActiveWindow.View
            view.ShowDrawings = True
            view.ShowPicturePlaceHolders = False
            view.Type = 3
        except Exception:
            pass
        try:
            self._review_app.ScreenRefresh()
        except Exception:
            pass

    def close_review_document(self, save_changes: bool = False) -> None:
        review_document = self._review_document
        review_app = self._review_app
        com_initialized = self._review_com_initialized

        self._review_document = None
        self._review_app = None
        self._review_com_initialized = False

        try:
            if review_document is not None:
                review_document.Close(SaveChanges=save_changes)
        except Exception:
            pass
        try:
            if review_app is not None:
                review_app.Quit(SaveChanges=save_changes)
        except Exception:
            pass
        try:
            if com_initialized and pythoncom is not None:
                pythoncom.CoUninitialize()
        except Exception:
            pass

    def release_to_user(self) -> None:
        self._review_document = None
        self._review_app = None
        self._review_com_initialized = False

    def export_pdf(self, output_path: str) -> None:  # pragma: no cover
        raise RuntimeError("La exportación PDF no está implementada en el flujo headless.")

    def build_document_from_template(
        self,
        template_path: str,
        output_path: str,
        image_paths: list[Path],
        placeholder_count: int,
    ) -> None:
        pythoncom_module, win32_module = _load_pywin32()
        output_file = Path(output_path).resolve()
        output_file.parent.mkdir(parents=True, exist_ok=True)

        app = None
        document = None
        embedded = 0
        pythoncom_module.CoInitialize()
        try:
            app = win32_module.dynamic.Dispatch("Word.Application")
            app.Visible = False
            app.DisplayAlerts = 0
            document = app.Documents.Open(
                str(Path(template_path).resolve()),
                ConfirmConversions=False,
                ReadOnly=False,
                AddToRecentFiles=False,
                Visible=False,
            )

            for index in range(1, self.MAX_PLACEHOLDER_IMAGES + 1):
                placeholder = self._placeholder_text(index)
                image_path = image_paths[index - 1] if index <= placeholder_count and index - 1 < len(image_paths) else None
                replaced = self._replace_placeholder_with_picture_com(document, placeholder, image_path)
                if replaced:
                    embedded += 1

            try:
                document.SaveAs(str(output_file))
            except Exception:
                document.SaveAs2(str(output_file))
        finally:
            try:
                if document is not None:
                    document.Close(SaveChanges=False)
            except Exception:
                pass
            try:
                if app is not None:
                    app.Quit(SaveChanges=False)
            except Exception:
                pass
            try:
                pythoncom_module.CoUninitialize()
            except Exception:
                pass

        self._embedded_image_count = embedded

    def build_document_without_com(
        self,
        template_path: str,
        output_path: str,
        image_paths: list[Path],
        placeholder_count: int,
    ) -> None:
        self._ensure_docx_support()
        template_file = Path(template_path).resolve()
        output_file = Path(output_path).resolve()
        if not template_file.exists():
            raise FileNotFoundError(f"La plantilla Word no existe: {template_file}")

        doc = Document(str(template_file))
        self._replace_placeholders_without_com(doc, image_paths, placeholder_count)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))

    def _replace_placeholders_without_com(self, doc, image_paths: list[Path], placeholder_count: int) -> None:
        image_map = {
            self._placeholder_text(index): Path(image_paths[index - 1]) if index - 1 < len(image_paths) else None
            for index in range(1, placeholder_count + 1)
        }
        for placeholder_number in range(placeholder_count + 1, self.MAX_PLACEHOLDER_IMAGES + 1):
            image_map[self._placeholder_text(placeholder_number)] = None

        embedded_count = 0
        paragraphs = list(self._iter_docx_paragraphs(doc))
        for paragraph in reversed(paragraphs):
            original = paragraph.text
            if "<img" not in original:
                continue
            embedded_count += self._replace_placeholders_from_map(paragraph, image_map, doc)

        self._embedded_image_count = embedded_count

    def _replace_placeholders_from_map(self, paragraph, image_map: dict[str, Path | None], doc=None) -> int:
        placeholder_pattern = re.compile(r"<img(\d+)>")
        text = paragraph.text
        matches = list(placeholder_pattern.finditer(text))
        if not matches:
            return 0

        owner_doc = doc or getattr(paragraph.part, "document", None)
        target_width, target_height = self._resolve_template_content_size(owner_doc)
        target_paragraph = self._resolve_placeholder_target_paragraph(paragraph)
        if WD_ALIGN_PARAGRAPH is not None:
            target_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        target_paragraph.paragraph_format.space_before = 0
        target_paragraph.paragraph_format.space_after = 0
        if target_paragraph is not paragraph:
            paragraph.clear()
        target_paragraph.clear()
        inserted = 0
        cursor = 0
        for match in matches:
            if match.start() > cursor:
                target_paragraph.add_run(text[cursor:match.start()])

            placeholder = match.group(0)
            image_path = image_map.get(placeholder)
            if image_path is not None and image_path.exists():
                run = target_paragraph.add_run()
                run.add_picture(
                    str(image_path),
                    width=target_width,
                    height=target_height,
                )
                inserted += 1

            cursor = match.end()

        if cursor < len(text):
            target_paragraph.add_run(text[cursor:])

        return inserted

    def _replace_placeholder_in_paragraph(self, paragraph, *, placeholder: str, image_path: Path | None) -> bool:
        return bool(self._replace_placeholders_from_map(paragraph, {placeholder: image_path}))

    def _iter_docx_paragraphs(self, doc):
        yield from doc.paragraphs
        for table in doc.tables:
            yield from self._iter_table_paragraphs(table)
        for section in doc.sections:
            yield from section.header.paragraphs
            for table in section.header.tables:
                yield from self._iter_table_paragraphs(table)
            yield from section.footer.paragraphs
            for table in section.footer.tables:
                yield from self._iter_table_paragraphs(table)

    def _iter_table_paragraphs(self, table):
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested_table in cell.tables:
                    yield from self._iter_table_paragraphs(nested_table)

    def _replace_all_exact_text(self, text: str, replacement: str) -> bool:
        replaced = False
        for paragraph in self._iter_docx_paragraphs(self.document):
            if text not in paragraph.text:
                continue
            paragraph.text = paragraph.text.replace(text, replacement)
            replaced = True
        return replaced

    def _contains_exact_text(self, text: str) -> bool:
        return any(text in paragraph.text for paragraph in self._iter_docx_paragraphs(self.document))

    def _placeholder_text(self, placeholder_number: int) -> str:
        return f"<img{placeholder_number}>"

    @staticmethod
    def _ensure_docx_support() -> None:
        if Document is None or Cm is None:
            raise RuntimeError("python-docx no está instalado. Instalá las dependencias desde requirements.txt.")

    def _replace_placeholder_with_picture_com(self, document, placeholder: str, image_path: Path | None) -> bool:
        search_range = document.Content
        find = search_range.Find
        find.ClearFormatting()
        find.Text = placeholder
        find.Forward = True
        find.Wrap = 0
        find.MatchCase = True
        find.MatchWholeWord = False
        find.MatchWildcards = False
        if not find.Execute():
            return False

        target_range = search_range.Duplicate
        target_range.Text = ""
        if image_path is None or not image_path.exists():
            return False

        try:
            target_range.ParagraphFormat.Alignment = 1
            target_range.ParagraphFormat.SpaceBefore = 0
            target_range.ParagraphFormat.SpaceAfter = 0
        except Exception:
            pass

        inline_shape = document.InlineShapes.AddPicture(
            FileName=str(image_path.resolve()),
            LinkToFile=False,
            SaveWithDocument=True,
            Range=target_range,
        )
        inline_shape.LockAspectRatio = 0
        width_pt, height_pt = self._resolve_template_content_size_com(document)
        inline_shape.Width = width_pt
        inline_shape.Height = height_pt
        return True

    def _resolve_template_content_size(self, doc):
        if doc is None:
            return Cm(WORD_LABEL_WIDTH_CM), Cm(WORD_LABEL_HEIGHT_CM)
        return Cm(WORD_LABEL_WIDTH_CM), Cm(WORD_LABEL_HEIGHT_CM)

    def _resolve_placeholder_target_paragraph(self, paragraph):
        if Paragraph is None:
            return paragraph

        candidate = paragraph
        sibling = paragraph._element.getprevious()
        while sibling is not None and sibling.tag.endswith("}p"):
            previous = Paragraph(sibling, paragraph._parent)
            if previous.text.strip():
                break
            candidate = previous
            sibling = sibling.getprevious()
        return candidate

    def _resolve_template_content_size_com(self, document) -> tuple[float, float]:
        cm_to_points = 72.0 / 2.54
        width = WORD_LABEL_WIDTH_CM * cm_to_points
        height = WORD_LABEL_HEIGHT_CM * cm_to_points
        return width, height
