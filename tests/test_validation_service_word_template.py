from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from services.validation_service import ValidationError, ValidationService


class ValidationServiceWordTemplateTests(unittest.TestCase):
    def test_validate_word_template_placeholders_accepts_template_with_img_slots(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "template.docx"
            doc = Document()
            doc.add_paragraph(" ".join(f"<img{index}>" for index in range(1, 28)))
            doc.save(path)

            ValidationService().validate_word_template_placeholders(path)

    def test_validate_word_template_placeholders_rejects_template_without_img_slots(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "template.docx"
            doc = Document()
            doc.add_paragraph("No placeholders here")
            doc.save(path)

            with self.assertRaises(ValidationError):
                ValidationService().validate_word_template_placeholders(path)


if __name__ == "__main__":
    unittest.main()
