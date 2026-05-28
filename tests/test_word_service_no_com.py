from __future__ import annotations

import inspect
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from PIL import Image
from docx import Document as DocxDocument


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import services.word_service as word_service_module
from services.word_service import WordService


class WordServiceNoComTests(unittest.TestCase):
    def test_build_document_without_com_replaces_placeholders(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"
            image_one = temp_path / "image-1.png"
            image_two = temp_path / "image-2.png"

            self._write_image(image_one, (220, 40, 40))
            self._write_image(image_two, (40, 110, 220))

            template = DocxDocument()
            template.add_paragraph("Body <img1>")
            table = template.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Table <img2>"
            template.sections[0].header.paragraphs[0].text = "Header <img1>"
            template.sections[0].footer.paragraphs[0].text = "Footer <img2>"
            template.save(template_path)

            service = WordService()
            service.build_document_without_com(
                str(template_path),
                str(output_path),
                [image_one, image_two],
                2,
            )

            self.assertTrue(output_path.exists())
            self.assertEqual(4, service.count_visible_images())

            generated = DocxDocument(output_path)
            collected_text = self._collect_text(generated)
            self.assertIn("Body", collected_text)
            self.assertIn("Table", collected_text)
            self.assertIn("Header", collected_text)
            self.assertIn("Footer", collected_text)
            self.assertNotIn("<img1>", collected_text)
            self.assertNotIn("<img2>", collected_text)

            with zipfile.ZipFile(output_path) as archive:
                media = [name for name in archive.namelist() if name.startswith("word/media/")]
                self.assertGreaterEqual(len(media), 2)

    def test_replace_image_placeholder_blanks_missing_image_without_raising(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            template = DocxDocument()
            template.add_paragraph("<img1>")
            template.save(template_path)

            service = WordService()
            service.open(str(template_path))
            result = service.replace_image_placeholder(1, temp_path / "missing.png")

            self.assertFalse(result.adjusted)
            self.assertEqual([], service.find_remaining_image_placeholders())
            self.assertIn("imagen faltante", result.details)

    def test_build_document_without_com_uses_leading_blank_paragraph_for_placeholder_page(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"
            image_one = temp_path / "image-1.png"

            self._write_image(image_one, (220, 40, 40))

            template = DocxDocument()
            template.add_paragraph("")
            template.add_paragraph("")
            template.add_paragraph("<img1>")
            template.save(template_path)

            service = WordService()
            service.build_document_without_com(
                str(template_path),
                str(output_path),
                [image_one],
                1,
            )

            generated = DocxDocument(output_path)
            first_xml = generated.paragraphs[0]._element.xml
            third_text = generated.paragraphs[2].text
            self.assertIn("graphicData", first_xml)
            self.assertEqual("", third_text)

    def test_build_document_without_com_keeps_one_image_per_placeholder_group(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"
            images = [temp_path / f"image-{index}.png" for index in range(1, 4)]

            for index, image_path in enumerate(images, start=1):
                self._write_image(image_path, (30 * index, 40 * index, 50 * index))

            template = DocxDocument()
            for placeholder_number in range(1, 4):
                template.add_paragraph("")
                template.add_paragraph("")
                template.add_paragraph(f"<img{placeholder_number}>")
            template.save(template_path)

            service = WordService()
            service.build_document_without_com(
                str(template_path),
                str(output_path),
                images,
                3,
            )

            with zipfile.ZipFile(output_path) as archive:
                xml = archive.read("word/document.xml").decode("utf-8", "ignore")
                self.assertEqual(3, xml.count("<w:drawing"))
                self.assertNotIn("<img1>", xml)
                self.assertNotIn("<img2>", xml)
                self.assertNotIn("<img3>", xml)

    def test_word_service_source_keeps_generation_headless_and_review_isolated(self) -> None:
        source = inspect.getsource(word_service_module)
        self.assertIn("build_document_without_com", source)
        self.assertIn("run.add_picture(", source)
        self.assertIn("open_for_review", source)
        self.assertIn("Documents.Open", source)
        self.assertNotIn("PrintOut(", source)

    @staticmethod
    def _collect_text(doc: DocxDocument) -> str:
        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(paragraph.text for paragraph in cell.paragraphs)
        for section in doc.sections:
            parts.extend(paragraph.text for paragraph in section.header.paragraphs)
            parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
        return "\n".join(parts)

    @staticmethod
    def _write_image(path: Path, rgb: tuple[int, int, int]) -> None:
        image = Image.new("RGB", (64, 32), rgb)
        image.save(path)


if __name__ == "__main__":
    unittest.main()
